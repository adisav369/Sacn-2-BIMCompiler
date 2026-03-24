#!/usr/bin/env python3
"""Generate ANGKASA Technology Update Workshop DOCX from user-edited content."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Helper functions ──

def add_page_break():
    doc.add_page_break()

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            # Handle bold markers **text**
            parts = val.split('**')
            for pi, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    run.font.size = Pt(10)
                    if pi % 2 == 1:
                        run.bold = True
    doc.add_paragraph()  # spacing after table

def add_bold_para(text, size=11):
    """Add paragraph with **bold** markers parsed."""
    p = doc.add_paragraph()
    parts = text.split('**')
    for i, part in enumerate(parts):
        if part:
            run = p.add_run(part)
            run.font.size = Pt(size)
            if i % 2 == 1:
                run.bold = True
    return p

def add_blockquote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    parts = text.split('**')
    for i, part in enumerate(parts):
        if part:
            run = p.add_run(part)
            run.font.size = Pt(11)
            if i % 2 == 1:
                run.bold = True
    return p

def add_ascii_art(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    return p

def heading(level, text):
    h = doc.add_heading(text, level=level)
    return h


# ════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════

doc.add_paragraph()  # spacer
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('1-Day Technology Update Workshop')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AI Disruption and the Future of\nCooperative Enterprise Systems')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()
doc.add_paragraph()

# Metadata block
meta_items = [
    ('Prepared for:', 'ANGKASA IT Division (Bahagian Teknologi Maklumat)'),
    ('Prepared by:', 'red1'),
    ('Facilitators:', 'NUSA TEAM'),
    ('Target Date:', 'April 2026'),
    ('Classification:', 'Discussion Draft \u2014 For Internal Review'),
]
for label, value in meta_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + ' ')
    run.bold = True
    run.font.size = Pt(12)
    run = p.add_run(value)
    run.font.size = Pt(12)

add_page_break()

# ════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════

heading(1, 'Executive Summary')

add_bold_para(
    'ANGKASA sits at a critical juncture. Its core SPGA salary deduction system \u2014 serving '
    '**3.5 million cooperative members** across **16,284 cooperatives** \u2014 runs on '
    '**Unisys ClearPath technology deployed since 1996**. This 30-year-old mainframe backbone '
    'has been reliable, but the world around it is changing at unprecedented speed.'
)

add_bold_para(
    'This paper presents the case that **any ERP procurement exercise conducted in 2026 '
    'without accounting for AI disruption risks locking ANGKASA into another generation of '
    'vendor dependency** \u2014 precisely at the moment when the economics of enterprise software '
    'are being fundamentally rewritten.'
)

add_bold_para(
    'The opportunity is equally significant: ANGKASA\u2019s unique position as the '
    '**apex body for 16,284 cooperatives with RM165.8 billion in assets** makes it an ideal '
    'incubator for a **cooperative-owned technology platform** that could serve the entire '
    'movement \u2014 at a fraction of proprietary costs.'
)

add_page_break()

# ════════════════════════════════════════════════════════════════
# PART 1: THE FUD THAT SHOULD CONCERN YOU
# ════════════════════════════════════════════════════════════════

heading(1, 'Part 1: The FUD That Should Concern You')

heading(2, '1.1 Your ERP RFP May Already Be Outdated')

doc.add_paragraph(
    'If ANGKASA is evaluating enterprise systems using a traditional RFP framework \u2014 '
    'comparing SAP, Oracle, and Microsoft on features, pricing, and implementation timelines '
    '\u2014 the evaluation criteria themselves may be obsolete within 18 months.'
)

add_bold_para(
    '**Why?** Because the ERP market is undergoing its most significant disruption in 35 years:'
)

add_blockquote(
    '\u201cThrough 2027, GenAI and AI agent use will create the first true challenge to '
    'mainstream productivity tools in 35 years, prompting a $58 billion market shake-up.\u201d\n'
    '\u2014 Gartner, 2025'
)

add_bold_para('**Key data points your RFP evaluation may not have considered:**')

add_table(
    ['Indicator', '2024', '2026 (Now)', '2027 Projection'],
    [
        ['Enterprise apps with AI agents', '<5%', '~20%', '**40%** (Gartner)'],
        ['ERP spending including AI capabilities', '14%', '~35%', '**62%** (Gartner)'],
        ['AI-native ERP startups funded', 'Emerging', 'DualEntry ($90M), others', '$270B AI software market'],
        ['SAP ECC end of support', 'Active', '22 months left', '**Dec 31, 2027**'],
    ]
)

add_bold_para(
    '**The risk:** Selecting a vendor today based on 2024-era criteria means you are buying a '
    'system designed for **the era that is ending**, not the one beginning.'
)

heading(2, '1.2 The SAP 2027 Cliff')

add_bold_para(
    'SAP has set a hard deadline: **mainstream ECC support ends December 31, 2027**. '
    'Extended maintenance (to 2030) costs a **2% premium** on top of the existing 22% '
    'maintenance fee. Organizations that wait face **30-50% higher consulting rates** due to '
    'the global migration demand surge.'
)

add_bold_para(
    '**What this means for ANGKASA:** Any vendor proposing SAP is proposing a system whose '
    'predecessor is being forcibly retired. The migration track record is sobering:'
)

add_bullet('**32% of companies** migrating to S/4HANA have **underestimated total cost**')
add_bullet('Typical migration: **12-24 months**, often extending beyond')
add_bullet('Organizations on unsupported ECC face **40-60% higher total cost of ownership**')

heading(2, '1.3 Oracle and Microsoft Are Raising Prices')

doc.add_paragraph(
    'The proprietary ERP vendors are coordinating price increases while bundling AI features '
    'that may or may not deliver value:'
)

add_table(
    ['Vendor', 'Current Pricing', 'Recent Changes'],
    [
        ['**Oracle Fusion Cloud ERP**', 'RM33,000/user/year ($7,500)', '3-year minimum contract, minimum 25 users'],
        ['**SAP S/4HANA Cloud**', 'RM31,000-66,000+/user/year', 'Forced migration from ECC by 2027'],
        ['**Microsoft Dynamics 365**', 'RM9,500-20,600/user/year', 'First price increase in 5 years (Nov 2025); M365 suite +8-23% effective (Jul 2026)'],
    ]
)

add_bold_para(
    '**At Oracle pricing, 100 ANGKASA staff = RM3.3 million/year in license fees alone. '
    'At SAP pricing, the same could exceed RM6.6 million/year.**'
)

doc.add_paragraph(
    'This is before implementation, customisation, training, and ongoing support.'
)

heading(2, '1.4 The AI-Native Disruption Is Real')

add_bold_para(
    'In October 2025, a startup called **DualEntry** emerged from stealth with **$90 million** '
    'in funding (Lightspeed, Khosla Ventures, Google Ventures) and a provocative claim: they '
    'can migrate organisations off legacy ERP systems **in 24 hours**.'
)

add_bold_para(
    'Their investor thesis was published under the title: **"The Death of the Legacy ERP"**.'
)

doc.add_paragraph('Key claims:')

add_bullet('First ERP built from the ground up with AI (not AI bolted onto legacy code)')
add_bullet('Already processed **$100 billion** in journal entries')
add_bullet('Signed **thousands** of global users in months')
add_bullet('Valued at **$415 million** in 15 months')

add_bold_para(
    '**Whether DualEntry succeeds is not the point.** The point is that **$100M+ in venture '
    'capital is betting that legacy ERP can be replaced by AI-native alternatives**. This is '
    'a structural market shift, not a trend.'
)

heading(2, '1.5 Forrester\u2019s Warning on Vendor Lock-In')

doc.add_paragraph(
    'Forrester\u2019s 2026 enterprise software predictions contain a specific warning:'
)

add_blockquote(
    '\u201cConsolidating to a primary platform concentrates risk and neuters negotiation '
    'leverage, making the next major software decision a bet on a single vendor\u2019s security '
    'posture, pricing model, and innovation capacity for the next decade.\u201d'
)

add_blockquote(
    '\u201cVendors are using AI to drive new revenue streams while tightening control... '
    'The era of experimentation is over. The era of monetisation has begun.\u201d'
)

add_bold_para(
    '**Translation:** The vendors offering AI features are using them as a justification to raise '
    'prices and deepen lock-in. An RFP that selects a proprietary vendor today is not just '
    'buying software \u2014 it is accepting a decade of pricing decisions made by that vendor.'
)

add_page_break()

# ════════════════════════════════════════════════════════════════
# PART 2: THE STRATEGIC POSITION
# ════════════════════════════════════════════════════════════════

heading(1, 'Part 2: The Strategic Position')

heading(2, '2.1 ERP Vendor Landscape \u2014 Cost vs AI Readiness')

doc.add_paragraph(
    'The following matrix positions major ERP options on two axes that matter for '
    'ANGKASA\u2019s decision:'
)

add_page_break()

# Matrix 1 - Cost vs AI Readiness
add_ascii_art(
    '                        AI READINESS\n'
    '                    HIGH\n'
    '                      |\n'
    '    SAP S/4HANA   |  DualEntry (AI-native)\n'
    '    RM31-66K/u/y  |  (startup, unproven)\n'
    '                      |\n'
    '    Oracle Fusion|\n'
    '    RM33K/u/y      |\n'
    '                      |\n'
    '    MS Dynamics   |.....OPPORTUNITY ZONE.....\n'
    '    RM10-21K/u/y  |  Open Source + AI Layer\n'
    '                      |  (iDempiere/ERPNext +\n'
    '                      |    local AI integration)\n'
    '                      |\n'
    '    Odoo Ent.       |\n'
    '    RM1.3-2K/u/y  |\n'
    '                      |\n'
    '    ERPNext          |  iDempiere\n'
    '    <RM50/u/y      |  RM0 license\n'
    '                      |\n'
    '                    LOW\n'
    '                      +---------------------------------->\n'
    '                HIGH          COST            LOW\n'
    '                COST                            COST'
)

doc.add_paragraph()

add_bold_para('**Reading the matrix:**')

add_bullet(
    '**Top-left quadrant (High Cost, High AI):** SAP and Oracle are investing heavily in AI, '
    'but at extreme cost. ANGKASA would pay RM3-7M/year for 100 users and still be locked '
    'into their roadmap.'
)
add_bullet(
    '**Bottom-right quadrant (Low Cost, Low AI):** Open source ERP (iDempiere, ERPNext) '
    'costs almost nothing in licensing but has no native AI today.'
)
add_bullet(
    '**The Opportunity Zone (center-right):** An open-source ERP base enhanced with a '
    'targeted AI layer \u2014 built or integrated locally \u2014 combines low cost with relevant AI '
    'capability. **This is where ANGKASA should position.**'
)

add_page_break()

heading(2, '2.2 Deployment Strategy \u2014 Implementation Risk vs Cooperative Fit')

doc.add_paragraph(
    'This matrix evaluates deployment approaches against the unique requirements of a '
    '**federated cooperative movement**:'
)

# Matrix 2 - Implementation Risk vs Cooperative Fit
add_ascii_art(
    '                    COOPERATIVE FIT\n'
    '                HIGH\n'
    '                      |\n'
    '                      |  Cooperative        iDempiere\n'
    '                      |  Tech Incubator   Shared Platform\n'
    '                      |  (ANGKASA-owned)  (OSGi modular)\n'
    '                      |\n'
    '                      |        ERPNext\n'
    '                      |        (shared instance)\n'
    '                      |\n'
    '                      |              Odoo Enterprise\n'
    '                      |\n'
    '                      |  DualEntry       MS Dynamics\n'
    '                      |  (fast but\n'
    '                      |    no coop        SAP/Oracle\n'
    '                      |    model)          (enterprise-only,\n'
    '                      |                      no federation)\n'
    '                      |\n'
    '                    LOW\n'
    '                      +---------------------------------->\n'
    '                HIGH     IMPLEMENTATION      LOW\n'
    '                RISK            RISK              RISK'
)

doc.add_paragraph()

add_bold_para('**Reading the matrix:**')

add_bullet(
    '**Bottom-left (High Risk, Low Fit):** SAP/Oracle are designed for single enterprises, '
    'not federated cooperative movements. Implementation risk is extreme (12-24 months, '
    'cost overruns).'
)
add_bullet(
    '**Top-right (Low Risk, High Fit):** A cooperative-owned technology incubator using '
    'open-source ERP scores highest. Proven in the credit union sector (Velera serves 4,000 '
    'credit unions; CU*Answers is cooperative-owned core processing).'
)
add_bullet(
    '**iDempiere** sits in the best position for ANGKASA specifically: Java/OSGi modularity '
    'means each cooperative type can have custom plugins, it matches ANGKASA\u2019s enterprise '
    'IT culture, and it has an active Malaysia community.'
)

heading(2, '2.3 The Cost Reality at ANGKASA\u2019s Scale')

add_bold_para(
    'ANGKASA is not a single company. It is the technology backbone for **16,284 cooperatives** '
    'and **7.22 million members**. At this scale, vendor pricing models produce staggering numbers:'
)

add_table(
    ['Scenario', 'Annual License Cost (est.)', '5-Year TCO (est.)'],
    [
        ['Oracle Fusion (500 core users)', '**RM16.5 million**', 'RM82+ million'],
        ['SAP S/4HANA Cloud (500 core users)', '**RM15-33 million**', 'RM75-165+ million'],
        ['Microsoft Dynamics 365 (500 core users)', '**RM4.8-10.3 million**', 'RM24-52 million'],
        ['Odoo Enterprise (500 users)', '**RM650K-1M**', 'RM3.3-5 million'],
        ['iDempiere (self-hosted, cooperative-owned)', '**RM0 license**', 'Infrastructure + support only'],
        ['ERPNext (self-hosted)', '**<RM25K**', 'Infrastructure + support only'],
    ]
)

add_bold_para(
    '**The question is not "Can ANGKASA afford SAP?" The question is: "Can 16,284 cooperatives '
    'justify sending RM80-160 million to a foreign vendor over 5 years when cooperative-owned '
    'alternatives exist?"**'
)

add_page_break()

# ════════════════════════════════════════════════════════════════
# PART 3: ANGKASA AS TECHNOLOGY INCUBATOR
# ════════════════════════════════════════════════════════════════

heading(1, 'Part 3: ANGKASA as Technology Incubator')

heading(2, '3.1 The Precedent: Credit Union Shared Services')

doc.add_paragraph(
    'The model already exists \u2014 not in theory, but in production at massive scale:'
)

add_table(
    ['Organisation', 'Model', 'Scale'],
    [
        ['**Velera** (formerly PSCU/Co-op Solutions)', 'Cooperative-owned payments CUSO', '**4,000+ credit unions** across North America'],
        ['**CU*Answers** (CUaxis/CBX)', 'Cooperative-owned core data processing', 'Hundreds of credit unions'],
        ['**Co-op Shared Branch Network**', 'Federated transaction network', '5,000+ branches'],
    ]
)

add_bold_para(
    '**These are cooperatives building technology platforms for cooperatives.** The members '
    'own the platform. The costs are shared. The vendor lock-in is zero.'
)

add_bold_para(
    'ANGKASA already operates this way for salary deductions (SPGA). The question is: '
    '**why not extend the model to ERP?**'
)

add_page_break()

heading(2, '3.2 The ANGKASA Incubator Concept')

# Incubator architecture diagram
add_ascii_art(
    '+-------------------------------------------------------------------+\n'
    '|                    ANGKASA TECHNOLOGY INCUBATOR                    |\n'
    '|                                                                   |\n'
    '|  +-------------------+  +-------------------+  +---------------+  |\n'
    '|  | Shared ERP        |  | AI Services       |  | Training &    |  |\n'
    '|  | Platform          |  | Layer             |  | Certification |  |\n'
    '|  | (iDempiere/       |  | (Local LLM,       |  | (HRD Corp     |  |\n'
    '|  |  open source)     |  |  document AI,     |  |  claimable)   |  |\n'
    '|  |                   |  |  workflow agents)  |  |               |  |\n'
    '|  +-------------------+  +-------------------+  +---------------+  |\n'
    '|          |                       |                      |         |\n'
    '|  +---------------------------------------------------------------+ |\n'
    '|  |             ANGKASA Data Center (Tier-3 Certified)            | |\n'
    '|  +---------------------------------------------------------------+ |\n'
    '|          |                       |                      |         |\n'
    '|  +--------+--------+  +--------+--------+  +--------+------+    |\n'
    '|  | Housing Coops   |  | Credit Coops    |  | Agri Coops    |    |\n'
    '|  | (BIM Compiler   |  | (Financial      |  | (SRI paddy,   |    |\n'
    '|  |  + BOM/cost)    |  |  management)    |  |  supply chain) |    |\n'
    '|  +--------+--------+  +--------+--------+  +--------+------+    |\n'
    '|  | Transport Coops |  | Consumer Coops  |  | Mosque Coops  |    |\n'
    '|  | (Fleet/fuel     |  | (Retail/        |  | (Waqf asset   |    |\n'
    '|  |  management)    |  |  inventory)     |  |  management)  |    |\n'
    '|  +-----------------+  +-----------------+  +---------------+    |\n'
    '+-------------------------------------------------------------------+\n'
    '                              |\n'
    '                   16,284 cooperatives\n'
    '                   7.22 million members\n'
    '                   RM165.8 billion assets'
)

doc.add_paragraph()

heading(2, '3.3 Why iDempiere Fits ANGKASA')

add_table(
    ['Requirement', 'iDempiere Capability'],
    [
        ['**Java enterprise stack**', 'Matches ANGKASA\u2019s existing IT culture (enterprise-grade, not startup-grade)'],
        ['**OSGi modularity**', 'Each cooperative type gets purpose-built plugins without forking the core'],
        ['**100% open source (GPLv2)**', 'Zero license fees. ANGKASA owns the code. No vendor can raise prices.'],
        ['**Active Malaysia community**', 'Redhuan D. Oon (author of Open Source ERP, Pearson Malaysia) is based in Putrajaya. Local expertise exists.'],
        ['**Manufacturing + Distribution**', 'Modules for supply chain, warehousing, POS \u2014 relevant to diverse cooperative types'],
        ['**Multi-organisation architecture**', 'Built-in support for multiple entities sharing one database \u2014 ideal for cooperative federation'],
        ['**20+ year track record**', 'Descended from Compiere (2000) \u2192 ADempiere (2006) \u2192 iDempiere (2012). Battle-tested.'],
    ]
)

heading(2, '3.4 The AI Layer \u2014 Built, Not Bought')

add_bold_para(
    'Rather than buying AI from SAP or Oracle (at their prices, on their terms), ANGKASA '
    'could **build an AI capability layer on top of open-source ERP**:'
)

add_table(
    ['AI Capability', 'Approach', 'Cooperative Use Case'],
    [
        ['**Document processing**', 'Local LLM (Ollama/vLLM) + OCR', 'Automate invoice/receipt processing for 16K cooperatives'],
        ['**Financial reporting agents**', 'Agentic AI on structured ERP data', 'Monthly close, audit preparation, regulatory compliance'],
        ['**BIM + Cost estimation**', 'BIM Intent Compiler (open source)', 'Housing cooperative design \u2192 BOM \u2192 cost estimate'],
        ['**Credit assessment**', 'AI on SPeKAR + cooperative financial data', 'Enhance loan assessment for cooperative members'],
        ['**Workflow automation**', 'AI agents replacing manual approval chains', 'Procurement, HR, membership management'],
    ]
)

add_bold_para(
    '**This is the "Opportunity Zone" from the strategy matrix \u2014 open source ERP + '
    'locally-built AI, owned by the cooperative movement.**'
)

add_page_break()

# ════════════════════════════════════════════════════════════════
# PART 4: THE DAKOM 2030 ALIGNMENT
# ════════════════════════════════════════════════════════════════

heading(1, 'Part 4: The DaKoM 2030 Alignment')

add_bold_para(
    'Malaysia\u2019s **Dasar Koperasi Malaysia (DaKoM) 2030** targets:'
)

add_bullet('Cooperative sector revenue of **RM73 billion** by 2030 (from RM45.5 billion)')
add_bullet('Active cooperative rate of **65%**')
add_bullet('Improvement in Cooperative Member Well-being Index (IndeKA) to **80**')

doc.add_paragraph('A technology incubator directly supports all three targets:')

add_table(
    ['DaKoM Target', 'Technology Contribution'],
    [
        ['**RM73B revenue**', 'Shared ERP reduces operational costs, freeing capital for growth. AI automation improves productivity.'],
        ['**65% active rate**', 'Low-cost shared platform makes digital tools accessible to small and inactive cooperatives currently priced out.'],
        ['**IndeKA score 80**', 'Better financial management, faster loan processing, transparent reporting improve member well-being.'],
    ]
)

add_bold_para(
    'Budget 2026 allocated **RM50 million** for cooperative development through SKM. A '
    'portion directed toward a cooperative technology incubator would have multiplied '
    'impact across all 16,284 cooperatives.'
)

add_page_break()

# ════════════════════════════════════════════════════════════════
# PART 5: WORKSHOP AGENDA
# ════════════════════════════════════════════════════════════════

heading(1, 'Part 5: Workshop Agenda')

heading(2, 'Morning \u2014 Understanding the Threat')

add_table(
    ['Time', 'Module', 'Content'],
    [
        ['09:00', '**Opening**', 'Workshop objectives. Why this matters now.'],
        ['09:15', '**Module 1: The AI Disruption**',
         'Gartner\u2019s 40% AI agent prediction. DualEntry and the "Death of Legacy ERP" thesis. '
         'What "AI-native" means vs "AI-bolted-on". Demo: AI agent processing a cooperative financial report.'],
        ['10:15', '**Break**', ''],
        ['10:30', '**Module 2: The Pricing Trap**',
         'SAP 2027 deadline. Oracle\u2019s RM33K/user/year. Microsoft\u2019s 23% effective increase. '
         'Forrester\u2019s lock-in warning. Cost projections at ANGKASA\u2019s scale. Interactive exercise: '
         'calculate your own 5-year TCO under three scenarios.'],
        ['11:30', '**Module 3: Your Current Position**',
         'ANGKASA\u2019s Unisys ClearPath \u2014 30 years of service. SPGA system strengths and risks. '
         'Where the gaps are. Facilitated discussion: what keeps the IT team up at night?'],
        ['12:30', '**Lunch**', ''],
    ]
)

heading(2, 'Afternoon \u2014 Seeing the Opportunity')

add_table(
    ['Time', 'Module', 'Content'],
    [
        ['13:30', '**Module 4: The Open Source Alternative**',
         'iDempiere, ERPNext, Odoo \u2014 live comparison. Total cost of ownership at cooperative scale. '
         'Case study: Velera (4,000 credit unions on shared platform).'],
        ['14:30', '**Module 5: ANGKASA as Incubator**',
         'The cooperative technology incubator model. Architecture walkthrough. How OSGi modularity '
         'serves 16,284 diverse cooperatives. DaKoM 2030 alignment. Interactive: mapping cooperative '
         'types to ERP modules.'],
        ['15:30', '**Break**', ''],
        ['15:45', '**Module 6: The AI Layer You Can Own**',
         'Local AI deployment (not cloud-dependent). Document processing demo. BIM Compiler demo '
         '(housing cooperative use case). Financial reporting agent concept.'],
        ['16:30', '**Module 7: Strategic Planning Exercise**',
         '9-Quadrant SWOT analysis (Prof. Wan Mohtar\u2019s methodology). Teams draft a 90-day action plan. '
         'Presentation and critique.'],
        ['17:15', '**Closing**', 'Summary of key decisions. Next steps. Commitment to action.'],
    ]
)

add_page_break()

# ════════════════════════════════════════════════════════════════
# PART 6: RECOMMENDED NEXT STEPS FOR ANGKASA
# ════════════════════════════════════════════════════════════════

heading(1, 'Part 6: Recommended Next Steps for ANGKASA')

add_table(
    ['Priority', 'Action', 'Timeline', 'Owner'],
    [
        ['**1**', 'Pause or reframe any active ERP RFP to include AI-readiness and cooperative-fit criteria',
         'Immediate', 'IT Division'],
        ['**2**', 'Commission a formal TCO comparison: proprietary vs open-source at ANGKASA scale',
         '30 days', 'IT Division + Nusa'],
        ['**3**', 'Pilot iDempiere with one cooperative type (e.g., credit cooperative) as proof of concept',
         '90 days', 'IT Division + Redhuan D. Oon'],
        ['**4**', 'Establish a Technology Incubator steering committee with representation from diverse cooperative types',
         '60 days', 'Prof. Wan Mohtar + ANGKASA board'],
        ['**5**', 'Present incubator concept to SKM for alignment with DaKoM 2030 and Budget 2026 allocation',
         'Q3 2026', 'Prof. Wan Mohtar + ANGKASA CEO'],
        ['**6**', 'Develop HRDC-claimable training curriculum for cooperative digital readiness',
         'Q3 2026', 'Nusa'],
    ]
)

add_page_break()

# ════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════

heading(1, 'References')

references = [
    'Gartner. "40% of Enterprise Apps Will Feature Task-Specific AI Agents by 2026." August 2025. gartner.com',
    'Gartner. "Strategic Predictions for 2026: How AI\'s Underestimated Influence Is Reshaping Business." gartner.com',
    'Forrester. "Predictions 2026: AI Agents Impact Enterprise Software." forrester.com',
    'Forrester. "AI means bigger margins and lock-in for enterprise software vendors." The Register. theregister.com',
    'Lightspeed Venture Partners. "The Death of the Legacy ERP: Why We Led DualEntry\'s Series A." lsvp.com',
    'DualEntry. "$90M Series A Funding Announcement." dualentry.com',
    'Sapphire Ventures. "2026 Outlook: 10 AI Predictions." sapphireventures.com',
    'Foundation Capital. "Where AI is headed in 2026." foundationcapital.com',
    'Kellton. "SAP ECC to S/4HANA Migration: 2027 Deadline Guide." kellton.com',
    'Top10ERP. "Oracle Fusion Cloud ERP Pricing Guide 2026." top10erp.org',
    'SAMExpert. "Microsoft 365 Price Increase July 2026." samexpert.com',
    'Mordor Intelligence. "Open Source ERP Market \u2014 USD 2.85B (2025) to USD 4.60B (2030)." mordorintelligence.com',
    'Unisys. "ANGKASA Selects Unisys Software to Enhance Client Experience." unisys.com',
    'Unisys. "ANGKASA Renews Contract for Unisys ClearPath." prnewswire.com',
    'The Star. "Over 16,000 cooperatives registered in Malaysia as of Dec 2024." thestar.com.my',
    'BERNAMA. "53 Years On, Angkasa Maintains Successful Track Record." bernama.com',
    'SKM. "National Co-operative Policy (DaKoM) 2030." skm.gov.my',
    'The Star. "Budget 2026: Angkasa hails RM50mil support for cooperative development." thestar.com.my',
    'Oon, Redhuan D. Open Source ERP. Pearson Malaysia, 2010. ISBN 978-967-349-022-6.',
    'iDempiere. "Community Powered Enterprise \u2014 Free Open Source ERP and CRM." idempiere.org',
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    run.font.size = Pt(10)
    # Check for italic book title
    if 'Open Source ERP.' in ref and 'Oon' in ref:
        before = ref.split('Open Source ERP.')[0]
        after = ref.split('Open Source ERP.')[1]
        run2 = p.add_run(before)
        run2.font.size = Pt(10)
        run3 = p.add_run('Open Source ERP.')
        run3.font.size = Pt(10)
        run3.italic = True
        run4 = p.add_run(after)
        run4.font.size = Pt(10)
    else:
        run2 = p.add_run(ref)
        run2.font.size = Pt(10)

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), 'ANGKASA-TechUpdate-Workshop.docx')
doc.save(output_path)
print(f'DOCX saved to: {output_path}')
