#!/usr/bin/env python3
"""Replace ASCII art Matrix 1 with a proper Word table quadrant."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document('/home/red1/bim-compiler/EntrussVentures/ANGKASA-TechUpdate-Workshop.docx')

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell border. Usage: set_cell_border(cell, top=("sz","6"), bottom=...) """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{attrs[0]}" w:space="0" w:color="{attrs[1]}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_cell_text(cell, lines, bold_lines=None, font_size=9):
    """Add multiple lines to a cell with optional bold for specific line indices."""
    if bold_lines is None:
        bold_lines = []
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for i, line in enumerate(lines):
        if i > 0:
            run = p.add_run('\n')
        run = p.add_run(line)
        run.font.size = Pt(font_size)
        run.font.name = 'Calibri'
        if i in bold_lines:
            run.bold = True

# Find the ASCII art paragraph (Matrix 1) and its index
target_idx = None
for i, p in enumerate(doc.paragraphs):
    if 'AI READINESS' in p.text and 'OPPORTUNITY ZONE' in p.text:
        target_idx = i
        break

if target_idx is None:
    print("ERROR: Could not find Matrix 1 paragraph")
    exit(1)

print(f"Found Matrix 1 at paragraph index {target_idx}")

# Get the parent element and the paragraph element
para_element = doc.paragraphs[target_idx]._element
parent = para_element.getparent()

# Create the replacement content:
# Y-axis label + 3x3 table (axis label col + 2 quadrant cols)

# First, build the table: 3 rows x 3 cols
# Row 0: header row with axis labels
# Row 1: HIGH AI row (top quadrants)
# Row 2: LOW AI row (bottom quadrants)
# Col 0: Y-axis labels
# Col 1: HIGH COST quadrant
# Col 2: LOW COST quadrant

table = doc.add_table(rows=4, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(2.5)
    row.cells[1].width = Cm(6.5)
    row.cells[2].width = Cm(6.5)

# ── Row 0: X-axis header ──
# Merge top-left corner
c00 = table.rows[0].cells[0]
c00.text = ''
c01 = table.rows[0].cells[1]
c01.text = ''
p = c01.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('HIGH COST')
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

c02 = table.rows[0].cells[2]
c02.text = ''
p = c02.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LOW COST')
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ── Row 1: HIGH AI READINESS ──
# Y-axis label
c10 = table.rows[1].cells[0]
c10.text = ''
c10.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p = c10.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('HIGH\nAI\nREADINESS')
run.bold = True
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Top-left: High Cost, High AI
c11 = table.rows[1].cells[1]
set_cell_shading(c11, 'FFF3E0')  # light orange - expensive
add_cell_text(c11, [
    'SAP S/4HANA',
    'RM31-66K/user/year',
    '',
    'Oracle Fusion',
    'RM33K/user/year',
    '',
    'DualEntry (AI-native)',
    '(startup, unproven)',
], bold_lines=[0, 3, 6], font_size=9)
c11.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Top-right: Low Cost, High AI — OPPORTUNITY ZONE
c12 = table.rows[1].cells[2]
set_cell_shading(c12, 'E8F5E9')  # light green - opportunity
c12.text = ''
p = c12.paragraphs[0]
p.paragraph_format.space_before = Pt(6)

run = p.add_run('OPPORTUNITY ZONE')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

run = p.add_run('\n\nOpen Source + AI Layer')
run.bold = True
run.font.size = Pt(9)

run = p.add_run('\n(iDempiere/ERPNext +\n local AI integration)')
run.font.size = Pt(9)

c12.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# ── Row 2: LOW AI READINESS ──
c20 = table.rows[2].cells[0]
c20.text = ''
c20.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p = c20.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LOW\nAI\nREADINESS')
run.bold = True
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Bottom-left: High Cost, Low AI
c21 = table.rows[2].cells[1]
set_cell_shading(c21, 'FFEBEE')  # light red - worst
add_cell_text(c21, [
    'MS Dynamics',
    'RM10-21K/user/year',
], bold_lines=[0], font_size=9)
c21.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Bottom-right: Low Cost, Low AI
c22 = table.rows[2].cells[2]
set_cell_shading(c22, 'F5F5F5')  # light grey
add_cell_text(c22, [
    'Odoo Enterprise',
    'RM1.3-2K/user/year',
    '',
    'ERPNext  <RM50/user/year',
    '',
    'iDempiere  RM0 license',
], bold_lines=[0, 3, 5], font_size=9)
c22.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# ── Row 3: Arrow row (X-axis) ──
c30 = table.rows[3].cells[0]
c30.text = ''
c31 = table.rows[3].cells[1]
c31.text = ''
# Merge bottom cells for the arrow label
c31_merge = c31.merge(table.rows[3].cells[2])
p = c31_merge.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('COST \u2192 decreasing \u2192')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# Set minimal borders on axis cells
for cell in [c00, c10, c20, c30]:
    set_cell_border(cell, top=("0", "FFFFFF"), bottom=("0", "FFFFFF"),
                    start=("0", "FFFFFF"), end=("4", "999999"))

# Remove default table style borders, use custom
table.style = 'Table Grid'

# Now move the table element to replace the ASCII art paragraph
table_element = table._tbl
# Remove the table from where Document.add_table put it (end of doc)
table_element.getparent().remove(table_element)

# Insert the table where the ASCII art paragraph was
parent.replace(para_element, table_element)

# Save
output = '/home/red1/bim-compiler/EntrussVentures/ANGKASA-TechUpdate-Workshop.docx'
doc.save(output)
print(f"Matrix 1 replaced with quadrant table. Saved to {output}")
