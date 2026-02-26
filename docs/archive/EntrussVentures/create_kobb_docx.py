#!/usr/bin/env python3
"""Generate KOBB Petrol Station Workshop DOCX from the markdown content."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Helpers ──

def pb():
    doc.add_page_break()

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def tbl(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            parts = val.split('**')
            for pi, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    run.font.size = Pt(10)
                    if pi % 2 == 1:
                        run.bold = True
    doc.add_paragraph()

def bp(text, size=11):
    p = doc.add_paragraph()
    parts = text.split('**')
    for i, part in enumerate(parts):
        if part:
            run = p.add_run(part)
            run.font.size = Pt(size)
            if i % 2 == 1:
                run.bold = True
    return p

def bq(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    parts = text.split('**')
    for i, part in enumerate(parts):
        if part:
            run = p.add_run(part)
            run.font.size = Pt(11)
            if i % 2 == 1:
                run.bold = True

def mono(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def heading(level, text):
    return doc.add_heading(text, level=level)


# ════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('1-Day IT Literacy Workshop\nfor Petrol Station Operators')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Understanding Your Numbers:\nSENTINEL, SETEL, E-Invoice,\nand Why Data Accuracy Matters')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
doc.add_paragraph()

meta = [
    ('Prepared for:', 'KOBB (Koperasi Pengusaha Stesen Minyak Bumiputera) Members'),
    ('Facilitated by:', 'Entruss Ventures Sdn. Bhd. (1186369-A)'),
    ('Workshop Lead:', 'Redhuan D. Oon \u2014 Subject Matter Expert, FOSS ERP & AI'),
    ('Patron:', 'Prof. Dr. Hj. Wan Mohtar Wan Yusoff \u2014 Chairman, KOBB; EVSB Expert'),
    ('Classification:', 'HRDC-Claimable | 1-Day Programme (09:00\u201317:00)'),
    ('Date:', 'February 2026 (Draft)'),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + ' ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(value)
    run.font.size = Pt(11)

pb()

# ════════════════════════════════════════════════════════════════
# WHY THIS WORKSHOP EXISTS
# ════════════════════════════════════════════════════════════════

heading(1, 'Why This Workshop Exists')

doc.add_paragraph(
    'Since the rollout of Malaysia\u2019s targeted fuel subsidy programme \u2014 diesel (June 2024) '
    'and RON95 BUDI95 (September 2025) \u2014 petrol station operators have been caught between '
    'multiple systems that do not talk to each other the way operators expect.'
)

bp(
    'The result: **daily mismatches in sales figures, subsidy claims, and e-invoice records** '
    'that operators cannot explain, cannot reconcile, and cannot resolve without understanding '
    'what is actually happening inside these systems.'
)

bp(
    'This workshop **does not teach IT**. It teaches **principles** \u2014 in plain language \u2014 '
    'so that operators can:'
)
bullet('Understand *why* their numbers don\u2019t match')
bullet('Know *what to demand* from vendors and system providers')
bullet('Protect themselves from financial losses caused by data errors they did not create')

pb()

# ════════════════════════════════════════════════════════════════
# THE PROBLEM IN PLAIN LANGUAGE
# ════════════════════════════════════════════════════════════════

heading(1, 'The Problem in Plain Language')

heading(2, 'What Operators Experience Daily')

mono(
    'Morning:  You open the station. Pumps are running. SENTINEL is recording sales.\n'
    '          SETEL is processing digital payments. E-invoice is generating tax documents.\n'
    '\n'
    'Evening:  You close the station. You look at THREE different numbers:\n'
    '          - Pump meter totals (physical)\n'
    '          - SENTINEL system totals (subsidy tracking)\n'
    '          - SETEL transaction totals (digital payment)\n'
    '\n'
    '          They don\u2019t match.\n'
    '\n'
    'Month-end: LHDN wants e-invoices. Credit notes. Debit notes.\n'
    '           The numbers STILL don\u2019t match.\n'
    '           You don\u2019t know whose number is right.\n'
    '           You don\u2019t know whose number is wrong.\n'
    '           You don\u2019t even know who to call.'
)

heading(2, 'What We Have Diagnosed')

doc.add_paragraph(
    'Through direct engagement with KOBB station operators and meetings with SETEL, '
    'Entruss Ventures has identified three root causes \u2014 none of which are the operator\u2019s fault:'
)

tbl(
    ['Root Cause', 'What Happens', 'Impact'],
    [
        ['**1. Live System Access Conflict**',
         'SENTINEL and SETEL access the same transaction data but at different moments in time. '
         'They are two separate systems reading a live database that is constantly changing.',
         'Daily totals diverge. Operators cannot reconcile.'],
        ['**2. Vendor Update Practice**',
         'The system vendors update records, prices, and configurations without coordinating '
         'a consistent cutoff point. Data changes mid-day, mid-transaction, mid-batch.',
         'Yesterday\u2019s report doesn\u2019t match today\u2019s report for the same day.'],
        ['**3. Incomplete UUID Disclosure**',
         'Every transaction in SETEL has a UUID \u2014 but operators are not given the full UUID. '
         'They get truncated or partial references.',
         'Operators cannot trace specific disputed transactions.'],
    ]
)

pb()

# ════════════════════════════════════════════════════════════════
# PART 1: THE FIVE PRINCIPLES
# ════════════════════════════════════════════════════════════════

heading(1, 'Part 1: The Five Principles Every Operator Must Know')

# ── Principle 1 ──
heading(2, 'Principle 1: Source of Truth')

bq('"If two systems disagree, one of them is wrong. You need to decide IN ADVANCE which one is the master."')

doc.add_paragraph(
    'In every business \u2014 from a nasi lemak stall to Petronas headquarters \u2014 there must be '
    'one agreed source of truth for each type of data.'
)

tbl(
    ['Data Type', 'Who Should Be Source of Truth?', 'Current Reality'],
    [
        ['Litres dispensed', 'Physical pump meter', 'Pump meter is correct (calibrated by KPDN)'],
        ['Subsidy-eligible transactions', 'SENTINEL', 'SENTINEL records at time of transaction'],
        ['Digital payment amount', 'SETEL', 'SETEL records at time of payment'],
        ['Tax document (e-invoice)', 'LHDN MyInvois', 'Generated after the fact \u2014 from which source?'],
    ]
)

bp(
    '**The problem:** There is **no agreed single source of truth** that SENTINEL, SETEL, '
    'and the e-invoice system all read from at the same moment.'
)

bp(
    '**What operators should demand:** A clear written answer from Petronas/vendors: '
    '**"Which system is the master record for daily sales reconciliation?"**'
)

# ── Principle 2 ──
heading(2, 'Principle 2: GIGO \u2014 Garbage In, Garbage Out')

bq('"A computer does not know if the data is wrong. It will process garbage perfectly and give you a perfect-looking garbage report."')

doc.add_paragraph(
    'The vendors\u2019 systems are not broken. They process data exactly as programmed. But if '
    'the input data is inconsistent \u2014 because of timing differences, retroactive updates, or '
    'configuration changes \u2014 the output will be consistently wrong.'
)

bp('**Examples operators will recognise:**')

tbl(
    ['Scenario', 'What Went In (Garbage)', 'What Came Out (Garbage)'],
    [
        ['Price updated at 11:47am, but pump recorded sale at 11:46am',
         'One system uses old price, another uses new price',
         'RM difference on one transaction \u2014 multiplied by hundreds per day'],
        ['SETEL batch closes at 11:59pm, SENTINEL closes at 12:00am',
         'Transactions near midnight appear in different days',
         'Daily totals don\u2019t match. Neither system is "wrong".'],
        ['Credit note issued for subsidy adjustment, dated 3 days later',
         'E-invoice shows original amount on Day 1, credit note on Day 4',
         'Monthly reconciliation is chaos \u2014 discrepancy for 3 days'],
    ]
)

bp(
    '**What operators should demand:** All systems must use the **same cutoff time** '
    '(e.g., 00:00:00 midnight) and the **same price table version** when processing the same transactions.'
)

pb()

# ── Principle 3 ──
heading(2, 'Principle 3: EODA Backup \u2014 Your Snapshot Source of Truth')

bq('"EODA = End Of Day Archive. Take a backup snapshot every night. That frozen snapshot is YOUR truth."')

bp('**What is EODA?**')

mono(
    'Every day at closing time (e.g., 11:55 PM):\n'
    '\n'
    '1. EXPORT your pump meter readings (physical count)\n'
    '2. EXPORT your SENTINEL daily summary (if available)\n'
    '3. EXPORT your SETEL daily transaction list\n'
    '4. SAVE all three into one folder: EODA_2026-02-18/\n'
    '\n'
    'That folder is now your FROZEN SNAPSHOT for that date.\n'
    'No one can change it. No vendor update can alter it.\n'
    'It is YOUR record.'
)

bp('**Why this matters:**')

doc.add_paragraph(
    'When SETEL or SENTINEL\u2019s live system is updated retroactively (price corrections, '
    'late credit notes, subsidy adjustments), the live reports will change. But your EODA '
    'snapshot will not. You can now compare:'
)
bullet('What the system said on the night of February 18')
bullet('What the system says today about February 18')

bp(
    '**If they differ, you have proof that the vendor changed the data after the fact.** '
    'Without EODA, you have no proof.'
)

bp('**Recommendation for SETEL/SENTINEL vendors:**')

bq(
    'The subsidy reconciliation system should read from an EODA snapshot (frozen daily archive) '
    'rather than querying the live transactional database. Two separate systems accessing a live '
    'database at different times will NEVER produce identical numbers.'
)

# ── Principle 4 ──
heading(2, 'Principle 4: UUID \u2014 Your Transaction Fingerprint')

bq('"Every digital transaction has a UUID \u2014 a unique code that identifies it forever. If you don\u2019t have it, you cannot trace it. If you cannot trace it, you cannot dispute it."')

bp('**What is a UUID?**')

mono(
    'Example: 550e8400-e29b-41d4-a716-446655440000\n'
    '\n'
    'This is a 36-character code that uniquely identifies ONE transaction.\n'
    'No two transactions in the world will ever have the same UUID.\n'
    'It is the transaction\u2019s "IC number".'
)

bp('**The problem we uncovered:**')

doc.add_paragraph(
    'In our meeting with SETEL, we discovered that station operators are not given the full UUID '
    'for their transactions. They receive truncated references, partial codes, or internal batch '
    'numbers that cannot be traced back to the specific transaction in SETEL\u2019s system.'
)

bp('**Why this is unacceptable:**')

tbl(
    ['Scenario', 'With Full UUID', 'Without Full UUID'],
    [
        ['Customer disputes a charge',
         'Operator shows UUID \u2192 SETEL confirms exact transaction',
         'Operator has nothing to show. SETEL says "which transaction?"'],
        ['Subsidy mismatch on one sale',
         'Operator traces UUID in SENTINEL and SETEL \u2192 finds discrepancy',
         'Operator sees discrepancy but cannot identify which transaction'],
        ['LHDN audit on e-invoice',
         'Operator maps UUID to e-invoice line item \u2192 complete trail',
         'Operator cannot prove which sale corresponds to which invoice'],
        ['Credit note issued late',
         'UUID links original transaction \u2192 credit note \u2192 revised amount',
         'No linkage. Operator\u2019s books have unexplained gaps.'],
    ]
)

bp(
    '**What operators should demand:** Full UUID for every transaction, printed on receipts '
    'AND available in daily export files.'
)

pb()

# ── Principle 5 ──
heading(2, 'Principle 5: Reconciliation Must Be Tri-Party')

bq('"You cannot reconcile two systems by looking at only one of them."')

doc.add_paragraph(
    'Current practice expects operators to reconcile SENTINEL and SETEL separately. But '
    'subsidy transactions involve three parties:'
)

mono(
    '          PUMP (Physical)\n'
    '              |\n'
    '    +---------+---------+\n'
    '    |                   |\n'
    'SENTINEL            SETEL\n'
    '(Subsidy tracking)  (Payment processing)\n'
    '    |                   |\n'
    '    +---------+---------+\n'
    '              |\n'
    '         E-INVOICE\n'
    '         (LHDN tax document)'
)

bp('**All four numbers must agree for a given day:**')
bullet('Pump meters: X litres dispensed')
bullet('SENTINEL: X litres recorded as subsidised / non-subsidised')
bullet('SETEL: RM Y collected via digital payment')
bullet('E-Invoice: RM Y invoiced to LHDN')

bp(
    '**If ANY of these numbers disagree, there is a data integrity problem \u2014 '
    'and the operator should NOT sign off until it is resolved.**'
)

pb()

# ════════════════════════════════════════════════════════════════
# PART 2: THE E-INVOICE NIGHTMARE
# ════════════════════════════════════════════════════════════════

heading(1, 'Part 2: The E-Invoice Nightmare')

heading(2, 'What Operators Face')

doc.add_paragraph(
    'Since August 2024 (Phase 1) and continuing into 2025-2026, petrol station operators '
    'must comply with LHDN\u2019s mandatory e-invoicing requirements.'
)

bp('**The specific pain points for station operators:**')

tbl(
    ['Issue', 'What Happens', 'Impact'],
    [
        ['**Late credit notes**',
         'Subsidy adjustments generate credit notes days after the original sale',
         'Daily e-invoice total doesn\u2019t match daily cash/payment total.'],
        ['**Debit notes for underclaims**',
         'When the subsidy calculation is corrected upward, a debit note is issued',
         'Monthly reconciliation shows unexplained increases.'],
        ['**Timing mismatch**',
         'E-invoice is generated at a different time than the SENTINEL/SETEL record',
         'The same sale appears on different dates in different systems.'],
        ['**Batch vs real-time**',
         'E-invoices are generated in batches, but transactions are real-time',
         'Individual amounts may be aggregated differently.'],
    ]
)

heading(2, 'The Root Cause')

doc.add_paragraph(
    'The e-invoice system was designed for simple buy-sell transactions: one seller, one buyer, '
    'one invoice, one payment. The petrol subsidy system introduces a third party (the government '
    'subsidy), which means:'
)

bullet('The **pump price** (what the customer sees): RM2.05/litre (RON95)')
bullet('The **subsidised price** (what the customer pays): RM1.99/litre')
bullet('The **subsidy amount** (what the government pays): RM0.06/litre')

bp(
    '**Three amounts for one transaction. One e-invoice. How do you reconcile?**'
)

doc.add_paragraph(
    'This is not an operator problem. This is a system design problem that LHDN, KPDN, '
    'and the fuel companies need to resolve. But operators bear the consequences.'
)

heading(2, 'What Operators Should Do Now')

bullet('**Keep EODA snapshots** \u2014 your daily frozen record (Principle 3)')
bullet('**Demand full UUIDs** \u2014 so you can trace any disputed transaction (Principle 4)')
bullet('**Log every credit/debit note** with the date it was issued AND the date of the original transaction')
bullet('**Do not sign off on monthly reconciliations** where the gap exceeds a threshold you are comfortable with')
bullet('**Report discrepancies formally** to KOBB \u2014 aggregate data strengthens the cooperative\u2019s negotiating position')

pb()

# ════════════════════════════════════════════════════════════════
# PART 3: WHAT KOBB CAN DO
# ════════════════════════════════════════════════════════════════

heading(1, 'Part 3: What KOBB Can Do as a Cooperative')

heading(2, 'Individual Operators Are Weak. The Cooperative Is Strong.')

doc.add_paragraph(
    'A single petrol station operator asking SETEL for full UUIDs will be ignored. '
    'KOBB representing all its member stations asking for full UUIDs is a formal demand '
    'from a registered cooperative with institutional backing.'
)

heading(2, 'KOBB\u2019s Collective Leverage')

tbl(
    ['Action', 'Individual Operator', 'KOBB Cooperative'],
    [
        ['Demand full UUID disclosure', 'Ignored', 'Formal request backed by SKM/ANGKASA'],
        ['Demand EODA-based reconciliation', '"That\u2019s not how the system works"', 'Negotiate as standard practice for all member stations'],
        ['Report persistent mismatches', 'One complaint among thousands', 'Aggregated data showing systemic pattern'],
        ['Engage LHDN on e-invoice timing', 'Complex and intimidating', 'KOBB submits formal industry feedback'],
        ['Demand vendor SLA for credit/debit note timing', 'No leverage', 'Contractual requirement for all member stations'],
        ['Training and capacity building', 'Self-funded, ad hoc', 'HRDC-claimable, structured, ongoing'],
    ]
)

heading(2, 'Prof. Wan Mohtar\u2019s Role')

doc.add_paragraph(
    'As Chairman of KOBB and board member of KOPSYA with strong ANGKASA connections, '
    'Prof. Dr. Hj. Wan Mohtar Wan Yusoff provides the institutional bridge between:'
)
bullet('**Operators on the ground** (who experience the problems daily)')
bullet('**ANGKASA** (which represents the cooperative movement nationally)')
bullet('**Regulators** (KPDN, LHDN, MOF) who set the rules')
bullet('**Vendors** (Petronas, SETEL) who build and operate the systems')

bp('This workshop is the first step in **converting individual frustration into collective action**.')

pb()

# ════════════════════════════════════════════════════════════════
# PART 4: WORKSHOP AGENDA
# ════════════════════════════════════════════════════════════════

heading(1, 'Part 4: Workshop Agenda')

heading(2, 'Morning \u2014 Understanding Your Systems')

tbl(
    ['Time', 'Module', 'Content', 'Method'],
    [
        ['09:00', '**Opening**', 'Welcome by Prof. Wan Mohtar. Workshop objectives.', 'Address'],
        ['09:15', '**Module 1: Source of Truth**', 'What it means. Why your three systems disagree. Interactive poll.', 'Lecture + poll'],
        ['09:45', '**Module 2: GIGO in Action**', 'Live walkthrough of how a single price update creates mismatched numbers.', 'Live demo'],
        ['10:15', '**Break**', '', ''],
        ['10:30', '**Module 3: SENTINEL vs SETEL**', 'What each system does. Why they see different numbers. Clock diagram.', 'Diagram + Q&A'],
        ['11:15', '**Module 4: The UUID Problem**', 'What a UUID is. Why you need the full one. What we found in our SETEL meeting.', 'Lecture + examples'],
        ['12:00', '**Module 5: EODA Backup**', 'Hands-on: how to take an EODA snapshot. Participants practice on their own data.', 'Hands-on exercise'],
        ['12:45', '**Lunch**', '', ''],
    ]
)

heading(2, 'Afternoon \u2014 Protecting Yourself')

tbl(
    ['Time', 'Module', 'Content', 'Method'],
    [
        ['13:45', '**Module 6: E-Invoice for Operators**', 'What LHDN requires. Credit notes, debit notes, timing. Three-party transactions.', 'Lecture + examples'],
        ['14:30', '**Module 7: Reconciliation Workshop**', 'Hands-on: reconcile one week of data. Identify and classify gaps by root cause.', 'Group exercise'],
        ['15:30', '**Break**', '', ''],
        ['15:45', '**Module 8: Your Rights as a Cooperative Member**', 'What KOBB can demand collectively. Template: formal discrepancy report form.', 'Discussion + template'],
        ['16:30', '**Module 9: Action Plan**', 'Personal action plan: Start EODA, demand UUID, log credit/debit notes, report to KOBB.', 'Individual exercise'],
        ['17:00', '**Closing**', 'Summary. Reference card distribution. Feedback forms. Certificate of attendance.', 'Address'],
    ]
)

pb()

# ════════════════════════════════════════════════════════════════
# PART 5: DELIVERABLES
# ════════════════════════════════════════════════════════════════

heading(1, 'Part 5: Deliverables to Participants')

doc.add_paragraph('Each participant receives:')

tbl(
    ['Item', 'Description'],
    [
        ['**Laminated Reference Card**', 'The 5 Principles on a single card \u2014 kept at the cash register for daily reference'],
        ['**EODA Template**', 'A simple folder structure and checklist for nightly backup snapshots'],
        ['**Discrepancy Report Template**', 'Standard form for reporting mismatches to KOBB \u2014 includes all relevant fields'],
        ['**UUID Demand Letter Template**', 'Draft letter from KOBB to SETEL requesting full UUID disclosure'],
        ['**Certificate of Attendance**', 'HRDC-claimable, issued by Entruss Ventures Sdn. Bhd.'],
    ]
)

# ════════════════════════════════════════════════════════════════
# PART 6: OUTCOME AND NEXT STEPS
# ════════════════════════════════════════════════════════════════

heading(1, 'Part 6: Outcome and Next Steps')

heading(2, 'For KOBB')

tbl(
    ['Action', 'Timeline', 'Lead'],
    [
        ['Circulate UUID demand letter to SETEL (via KOBB board)', 'Within 2 weeks of workshop', 'Prof. Wan Mohtar'],
        ['Collect first month of EODA snapshots from 10 pilot stations', '30 days', 'Workshop participants'],
        ['Aggregate discrepancy data and present to ANGKASA', '60 days', 'KOBB IT committee'],
        ['Submit formal industry feedback to LHDN on e-invoice timing', '90 days', 'KOBB + ANGKASA'],
        ['Plan follow-up workshop: "Advanced Reconciliation and Audit Preparation"', 'Q3 2026', 'Entruss Ventures'],
    ]
)

heading(2, 'For Entruss Ventures')

doc.add_paragraph(
    'This workshop establishes Entruss Ventures as a credible HRDC-certified training provider '
    'in the cooperative sector, specifically at the intersection of:'
)
bullet('**IT literacy for non-technical operators** \u2014 practical, not academic')
bullet('**Data governance and reconciliation** \u2014 principles that apply to any industry')
bullet('**Cooperative collective action** \u2014 translating individual problems into systemic solutions')

doc.add_paragraph(
    'The KOBB workshop is a replicable model that can be adapted for:'
)
bullet('Other petrol station cooperatives (non-Petronas brands)')
bullet('Agricultural cooperatives (similar subsidy tracking issues)')
bullet('Any cooperative sector facing multi-vendor system reconciliation')

pb()

# ════════════════════════════════════════════════════════════════
# APPENDIX A: GLOSSARY
# ════════════════════════════════════════════════════════════════

heading(1, 'Appendix A: Glossary for Operators')

tbl(
    ['Term', 'Malay', 'Plain English'],
    [
        ['**Source of Truth**', 'Sumber Kebenaran', 'The ONE system everyone agrees is the master record'],
        ['**GIGO**', 'Sampah Masuk, Sampah Keluar', 'Bad input data = bad output reports'],
        ['**EODA**', 'Arkib Akhir Hari', 'A frozen copy of your daily data, saved every night'],
        ['**UUID**', 'Nombor Pengenalan Transaksi', 'A unique code for every transaction \u2014 like an IC number'],
        ['**Reconciliation**', 'Penyesuaian', 'Comparing two or more sets of numbers to match'],
        ['**Credit Note**', 'Nota Kredit', 'A document that REDUCES the amount you owe'],
        ['**Debit Note**', 'Nota Debit', 'A document that INCREASES the amount you owe'],
        ['**E-Invoice**', 'E-Invois', 'Electronic invoice submitted to LHDN \u2014 mandatory'],
        ['**Snapshot**', 'Tangkap Layar Data', 'A frozen picture of data at one specific moment'],
        ['**Cutoff Time**', 'Masa Potong', 'The exact time when one day\u2019s data ends and the next begins'],
    ]
)

pb()

# ════════════════════════════════════════════════════════════════
# APPENDIX B: CLOCK DIAGRAM
# ════════════════════════════════════════════════════════════════

heading(1, 'Appendix B: The Clock Diagram \u2014 Why Systems Disagree')

mono(
    '         23:55         00:00          00:05\n'
    '           |             |              |\n'
    '    PUMP:  |===SALE===|  |              |     \u2190 Pump records sale at 23:57\n'
    '           |             |              |\n'
    'SENTINEL:  |    ......|BATCH CLOSE|    |     \u2190 SENTINEL closes batch at 00:00\n'
    '           |             |              |\n'
    '   SETEL:  |         |BATCH CLOSE|..   |     \u2190 SETEL closes batch at 00:02\n'
    '           |             |              |\n'
    '           |             |              |\n'
    '   RESULT: Pump says sale is on Feb 18\n'
    '           SENTINEL says sale is on Feb 18 (closed before midnight)\n'
    '           SETEL says sale is on Feb 19 (closed after midnight)\n'
    '\n'
    '           THREE systems. THREE dates. ONE sale.\n'
    '           NOBODY is wrong. The DESIGN is wrong.\n'
    '\n'
    '   FIX:    All systems must use the SAME cutoff time.\n'
    '           Or: all systems must read from the SAME EODA snapshot.'
)

doc.add_paragraph()
doc.add_paragraph()

# ── Footer ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Entruss Ventures Sdn. Bhd.')
run.bold = True
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    '(1186369-A)\n'
    'No. 15, Jalan 1/3c, Seksyen 1, 43650 Bandar Baru Bangi, Selangor D.E.\n'
    'Tel: 03-8925 8301 | Email: info@entruss.net'
)
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Workshop Enquiries: red1org@gmail.com\n')
run.font.size = Pt(10)
run = p.add_run('KOBB Cooperative: Prof. Dr. Hj. Wan Mohtar Wan Yusoff, Chairman')
run.font.size = Pt(10)

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), 'KOBB-PetrolStation-Workshop.docx')
doc.save(output_path)
print(f'DOCX saved to: {output_path}')
